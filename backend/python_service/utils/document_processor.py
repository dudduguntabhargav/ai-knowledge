import io
from typing import Tuple
from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_file(content: bytes, filename: str) -> Tuple[str, str]:
    file_extension = filename.lower().split('.')[-1]

    try:
        if file_extension == 'pdf':
            return extract_text_from_pdf(content), 'pdf'
        elif file_extension in ['docx', 'doc']:
            return extract_text_from_docx(content), 'docx'
        elif file_extension in ['txt', 'md', 'text']:
            return extract_text_from_txt(content), 'txt'
        else:
            return extract_text_from_txt(content), 'unknown'
    except Exception as e:
        raise ValueError(f"Failed to extract text from {filename}: {str(e)}")


def extract_text_from_pdf(content: bytes) -> str:
    pdf_file = io.BytesIO(content)
    pdf_reader = PdfReader(pdf_file)

    text_parts = []
    for page_num, page in enumerate(pdf_reader.pages, 1):
        page_text = page.extract_text()
        if page_text.strip():
            text_parts.append(f"--- Page {page_num} ---\n{page_text}")

    if not text_parts:
        raise ValueError("No text could be extracted from the PDF")

    return "\n\n".join(text_parts)


def extract_text_from_docx(content: bytes) -> str:
    docx_file = io.BytesIO(content)
    doc = Document(docx_file)

    text_parts = []

    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                text_parts.append(row_text)

    if not text_parts:
        raise ValueError("No text could be extracted from the DOCX file")

    return "\n\n".join(text_parts)


def extract_text_from_txt(content: bytes) -> str:
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("latin-1")
        except Exception:
            return content.decode("utf-8", errors="ignore")


def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list:
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        if end < len(text):
            last_period = text[end-100:end].rfind('.')
            last_newline = text[end-100:end].rfind('\n')

            boundary = max(last_period, last_newline)
            if boundary != -1:
                end = end - 100 + boundary + 1

        chunks.append(text[start:end].strip())
        start = end - overlap if end < len(text) else end

    return chunks
